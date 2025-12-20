"""
Knowledge Service - Document management and search for Knowledge Hub
"""
from datetime import datetime
from typing import List, Optional
from uuid import UUID, uuid4
import logging
from pathlib import Path
from sqlalchemy import text
import json
import io
import pdfplumber
import re

from fastapi import UploadFile
from sqlalchemy.orm import Session

from app.schemas.knowledge import (
    KnowledgeDocument,
    KnowledgeDocumentCreate,
    KnowledgeDocumentUpdate,
    KnowledgeDocumentList,
    KnowledgeDocumentUploadResponse,
    KnowledgeSearchRequest,
    KnowledgeSearchResponse,
    KnowledgeQueryRequest,
    KnowledgeQueryResponse,
)
from app.llm.gemini_client import GeminiChat, is_gemini_available
from app.llm.clients.jina_embed import embed_texts, is_jina_available
from app.vectorstore.pgvector_client import PgVectorClient
from app.services.storage_client import (
    build_object_key,
    generate_presigned_get_url,
    is_storage_configured,
    upload_bytes_to_storage,
    delete_object,
)
from app.core.config import get_settings

logger = logging.getLogger(__name__)

# In-memory storage for mock knowledge documents
_mock_knowledge_docs: dict[str, KnowledgeDocument] = {}


def _init_mock_knowledge_docs():
    """Initialize with mock knowledge documents"""
    mock_data = [
        {
            "id": UUID("a0000001-0000-0000-0000-000000000001"),
            "title": "Thông tư 09/2020/TT-NHNN - Quản lý rủi ro CNTT",
            "description": "Quy định về quản lý rủi ro công nghệ thông tin trong hoạt động ngân hàng",
            "document_type": "regulation",
            "source": "NHNN",
            "file_type": "pdf",
            "file_size": 2048000,
            "file_url": "/mock/knowledge/nhnn-09-2020.pdf",
            "tags": ["NHNN", "Compliance", "Risk Management", "IT"],
            "category": "Compliance",
            "uploaded_by": UUID("b0000001-0000-0000-0000-000000000001"),
            "uploaded_by_name": "Nguyễn Văn A",
            "uploaded_at": datetime.now(),
            "view_count": 45,
            "last_accessed_at": datetime.now(),
        },
        {
            "id": UUID("a0000002-0000-0000-0000-000000000002"),
            "title": "LPBank Security Policy v3.0",
            "description": "Chính sách bảo mật thông tin của LPBank - Phiên bản 3.0",
            "document_type": "policy",
            "source": "SharePoint",
            "file_type": "pdf",
            "file_size": 1536000,
            "file_url": "/mock/knowledge/security-policy-v3.pdf",
            "tags": ["Security", "Policy", "Encryption", "MFA"],
            "category": "Security",
            "uploaded_by": UUID("b0000004-0000-0000-0000-000000000004"),
            "uploaded_by_name": "Phạm Văn D - CTO",
            "uploaded_at": datetime.now(),
            "view_count": 32,
            "last_accessed_at": datetime.now(),
        },
        {
            "id": UUID("a0000003-0000-0000-0000-000000000003"),
            "title": "Core Banking Integration Guide",
            "description": "Hướng dẫn tích hợp với hệ thống Core Banking mới",
            "document_type": "technical",
            "source": "SharePoint",
            "file_type": "docx",
            "file_size": 3072000,
            "file_url": "/mock/knowledge/core-banking-integration.docx",
            "tags": ["Core Banking", "Integration", "API", "Technical"],
            "category": "Technical",
            "uploaded_by": UUID("b0000005-0000-0000-0000-000000000005"),
            "uploaded_by_name": "Hoàng Thị E - Tech Lead",
            "uploaded_at": datetime.now(),
            "view_count": 28,
            "last_accessed_at": datetime.now(),
        },
        {
            "id": UUID("a0000004-0000-0000-0000-000000000004"),
            "title": "KYC Policy 2024",
            "description": "Chính sách Know Your Customer - Cập nhật 2024",
            "document_type": "policy",
            "source": "LOffice",
            "file_type": "pdf",
            "file_size": 1024000,
            "file_url": "/mock/knowledge/kyc-policy-2024.pdf",
            "tags": ["KYC", "Policy", "Compliance", "Customer"],
            "category": "Compliance",
            "uploaded_by": UUID("b0000001-0000-0000-0000-000000000001"),
            "uploaded_by_name": "Nguyễn Văn A",
            "uploaded_at": datetime.now(),
            "view_count": 19,
            "last_accessed_at": datetime.now(),
        },
        {
            "id": UUID("a0000005-0000-0000-0000-000000000005"),
            "title": "Risk Assessment Template",
            "description": "Template đánh giá rủi ro dự án",
            "document_type": "template",
            "source": "SharePoint",
            "file_type": "xlsx",
            "file_size": 512000,
            "file_url": "/mock/knowledge/risk-assessment-template.xlsx",
            "tags": ["Risk", "Template", "Assessment"],
            "category": "Project",
            "uploaded_by": UUID("b0000009-0000-0000-0000-000000000009"),
            "uploaded_by_name": "Bùi Văn I - CRO",
            "uploaded_at": datetime.now(),
            "view_count": 15,
            "last_accessed_at": datetime.now(),
        },
        {
            "id": UUID("a0000006-0000-0000-0000-000000000006"),
            "title": "Thông tư 35/2016/TT-NHNN - Chuyển mạch thanh toán",
            "description": "Quy định về chuyển mạch thanh toán điện tử",
            "document_type": "regulation",
            "source": "NHNN",
            "file_type": "pdf",
            "file_size": 1792000,
            "file_url": "/mock/knowledge/nhnn-35-2016.pdf",
            "tags": ["NHNN", "Payment", "Compliance"],
            "category": "Compliance",
            "uploaded_by": UUID("b0000001-0000-0000-0000-000000000001"),
            "uploaded_by_name": "Nguyễn Văn A",
            "uploaded_at": datetime.now(),
            "view_count": 12,
            "last_accessed_at": datetime.now(),
        },
        {
            "id": UUID("a0000007-0000-0000-0000-000000000007"),
            "title": "Mobile Banking API Documentation",
            "description": "Tài liệu API cho ứng dụng Mobile Banking",
            "document_type": "technical",
            "source": "Wiki",
            "file_type": "md",
            "file_size": 768000,
            "file_url": "/mock/knowledge/mobile-api-docs.md",
            "tags": ["Mobile Banking", "API", "Documentation"],
            "category": "Technical",
            "uploaded_by": UUID("b0000006-0000-0000-0000-000000000006"),
            "uploaded_by_name": "Ngô Thị F - Tech Lead Mobile",
            "uploaded_at": datetime.now(),
            "view_count": 24,
            "last_accessed_at": datetime.now(),
        },
        {
            "id": UUID("a0000008-0000-0000-0000-000000000008"),
            "title": "Change Request Process Guide",
            "description": "Hướng dẫn quy trình xử lý Change Request",
            "document_type": "policy",
            "source": "SharePoint",
            "file_type": "pdf",
            "file_size": 1280000,
            "file_url": "/mock/knowledge/cr-process-guide.pdf",
            "tags": ["Change Request", "Process", "PMO"],
            "category": "Project",
            "uploaded_by": UUID("b0000001-0000-0000-0000-000000000001"),
            "uploaded_by_name": "Nguyễn Văn A",
            "uploaded_at": datetime.now(),
            "view_count": 18,
            "last_accessed_at": datetime.now(),
        },
    ]
    
    for doc_data in mock_data:
        doc = KnowledgeDocument(**doc_data)
        _mock_knowledge_docs[str(doc.id)] = doc


# Initialize mock data
_init_mock_knowledge_docs()


def _with_presigned_url(doc: KnowledgeDocument) -> KnowledgeDocument:
    """Attach a fresh presigned URL if the document is stored in S3."""
    if doc.storage_key and is_storage_configured():
        url = generate_presigned_get_url(doc.storage_key, expires_in=3600)
        if url:
            doc_dict = doc.model_dump()
            doc_dict["file_url"] = url
            return KnowledgeDocument(**doc_dict)
    return doc


def _row_to_doc(row) -> KnowledgeDocument:
    """Map DB row -> KnowledgeDocument; fill missing fields with defaults."""
    return KnowledgeDocument(
        id=row["id"],
        title=row["title"],
        description=row.get("description"),
        document_type=row.get("document_type") or "document",
        source=row.get("source") or "Uploaded",
        file_type=row.get("file_type") or "pdf",
        file_size=row.get("file_size"),
        file_url=row.get("file_url"),
        storage_key=row.get("storage_key"),
        tags=row.get("tags") or [],
        category=row.get("category"),
        uploaded_by=None,
        uploaded_by_name=None,
        uploaded_at=row.get("created_at") or datetime.now(),
        view_count=0,
        last_accessed_at=row.get("updated_at"),
    )


def _sanitize_text(text: str) -> str:
    """Remove NUL and trim."""
    if not text:
        return ""
    return text.replace("\x00", "").strip()


def _normalize_for_embedding(text: str) -> str:
    """Lowercase + sanitize for case-insensitive embeddings/search."""
    return _sanitize_text(text).lower()


def _is_smalltalk_or_noise(query: str) -> bool:
    """Heuristic: greetings or too-short queries => handle without RAG."""
    q = (query or "").strip().lower()
    if not q:
        return True
    # Very short or only punctuation
    if len(q) <= 2 or re.fullmatch(r"[.?/!]+", q):
        return True
    smalltalk_keywords = [
        "hi", "hello", "helo", "hey", "chào", "xin chào", "alo",
        "cái gì", "gì vậy", "gì đây", "ai đó", "bot ơi", "hello?", "test", "ping",
    ]
    return any(q == kw or q.startswith(kw + " ") for kw in smalltalk_keywords)


def _format_vector(vec: list[float]) -> str:
    """Format embedding list to Postgres vector literal."""
    return "[" + ",".join(f"{x:.6f}" for x in vec) + "]"


def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        return "\n".join(pages)
    except Exception as exc:
        logger.error("PDF extract failed: %s", exc)
        return ""


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx if available."""
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        logger.error("DOCX extract skipped (python-docx not available?): %s", exc)
        return ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: List[str] = []
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if text:
                parts.append(text)
        # Optionally extract table cells
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        logger.error("DOCX extract failed: %s", exc)
        return ""


def _extract_excel_text(file_bytes: bytes) -> str:
    """Extract text from XLSX/XLS using openpyxl if available."""
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception as exc:
        logger.error("Excel extract skipped (openpyxl not available?): %s", exc)
        return ""
    try:
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
        lines: List[str] = []
        for sheet in wb:
            lines.append(f"# {sheet.title}")
            for idx, row in enumerate(sheet.iter_rows(values_only=True)):
                if idx >= 200:  # avoid huge sheets
                    break
                cells = []
                for val in row:
                    if val is None:
                        continue
                    text = str(val).strip()
                    if text:
                        cells.append(text)
                if cells:
                    lines.append(" | ".join(cells))
            if len(lines) > 2000:
                break
        return "\n".join(lines)
    except Exception as exc:
        logger.error("Excel extract failed: %s", exc)
        return ""
def _chunk_text(text: str, max_len: int = 1200, overlap: int = 200) -> list[str]:
    """Greedy chunk by characters with overlap."""
    chunks = []
    start = 0
    n = len(text)
    while start < n:
        end = min(n, start + max_len)
        chunk = text[start:end]
        chunks.append(chunk)
        if end == n:
            break
        start = end - overlap
    return [c.strip() for c in chunks if c.strip()]


async def list_documents(
    db: Session,
    skip: int = 0,
    limit: int = 50,
    document_type: Optional[str] = None,
    source: Optional[str] = None,
    category: Optional[str] = None,
    meeting_id: Optional[UUID] = None,
    project_id: Optional[UUID] = None,
) -> KnowledgeDocumentList:
    """List all knowledge documents with optional filters"""
    try:
        conditions = ["1=1"]
        params = {"skip": skip, "limit": limit}
        if source:
            conditions.append("source = :source")
            params["source"] = source
        if category:
            conditions.append("category = :category")
            params["category"] = category
        if meeting_id:
            conditions.append("meeting_id = :meeting_id")
            params["meeting_id"] = str(meeting_id)
        if project_id:
            conditions.append("project_id = :project_id")
            params["project_id"] = str(project_id)

        where_clause = " AND ".join(conditions)
        rows = db.execute(
            text(
                f"""
                SELECT id, title, description, source, category, tags,
                       file_type, file_size, storage_key, file_url,
                       created_at, updated_at
                FROM knowledge_document
                WHERE {where_clause}
                ORDER BY created_at DESC NULLS LAST
                LIMIT :limit OFFSET :skip
                """
            ),
            params,
        ).mappings().all()

        total = db.execute(
            text(f"SELECT COUNT(*) FROM knowledge_document WHERE {where_clause}"),
            params,
        ).scalar_one()

        docs = [_with_presigned_url(_row_to_doc(r)) for r in rows]
        return KnowledgeDocumentList(documents=docs, total=total)
    except Exception as exc:
        logger.warning("List documents fallback to mock: %s", exc)
        docs = list(_mock_knowledge_docs.values())
        if document_type:
            docs = [d for d in docs if d.document_type == document_type]
        if source:
            docs = [d for d in docs if d.source == source]
        if category:
            docs = [d for d in docs if d.category == category]
        docs.sort(key=lambda x: x.uploaded_at, reverse=True)
        docs = [_with_presigned_url(d) for d in docs]
        return KnowledgeDocumentList(
            documents=docs[skip:skip + limit],
            total=len(docs),
        )


async def get_document(db: Session, document_id: UUID) -> Optional[KnowledgeDocument]:
    """Get a single document by ID and increment view count"""
    try:
        row = db.execute(
            text(
                """
                SELECT id, title, description, source, category, tags,
                       file_type, file_size, storage_key, file_url,
                       created_at, updated_at
                FROM knowledge_document
                WHERE id = :id
                """
            ),
            {"id": str(document_id)},
        ).mappings().first()
        if row:
            return _with_presigned_url(_row_to_doc(row))
    except Exception as exc:
        logger.warning("Get document fallback to mock: %s", exc)

    doc = _mock_knowledge_docs.get(str(document_id))
    if doc:
        doc_dict = doc.model_dump()
        doc_dict["view_count"] = doc.view_count + 1
        doc_dict["last_accessed_at"] = datetime.now()
        updated_doc = KnowledgeDocument(**doc_dict)
        _mock_knowledge_docs[str(document_id)] = updated_doc
        return _with_presigned_url(updated_doc)
    return None


async def upload_document(
    db: Session,
    data: KnowledgeDocumentCreate,
    file: Optional[UploadFile] = None,
) -> KnowledgeDocumentUploadResponse:
    """Upload a new knowledge document (mock metadata, real storage if configured)"""
    doc_id = uuid4()
    
    file_ext = data.file_type.lower()
    storage_key = None
    file_url = data.file_url or f"/mock/knowledge/{doc_id}.{file_ext}"
    file_size = data.file_size or 0

    # Handle actual file upload if provided
    if file:
        filename = file.filename or f"{doc_id}.{file_ext}"
        content = await file.read()
        file_size = len(content)

        if is_storage_configured():
            try:
                object_key = build_object_key(filename, prefix="knowledge")
                upload_bytes_to_storage(content, object_key, content_type=file.content_type)
                storage_key = object_key
                # Return a presigned URL so the frontend can access the private object
                presigned_url = generate_presigned_get_url(object_key, expires_in=3600)
                if presigned_url:
                    file_url = presigned_url
            except Exception as exc:
                logger.error("Upload to storage failed, falling back to local file: %s", exc)

        # Fallback to local storage if storage not configured or failed
        if not storage_key:
            upload_dir = Path(__file__).parent.parent / "uploaded_files"
            upload_dir.mkdir(parents=True, exist_ok=True)
            stored_name = f"{doc_id}.{file_ext}"
            stored_path = upload_dir / stored_name
            stored_path.write_bytes(content)
            file_url = f"/files/{stored_name}"

    # Get uploaded_by name (mock - would query user table in production)
    uploaded_by_name = "Current User"  # Would fetch from user table
    
    doc = KnowledgeDocument(
        id=doc_id,
        title=data.title,
        description=data.description,
        document_type=data.document_type,
        source=data.source or "Uploaded",
        file_type=data.file_type,
        file_size=file_size or 1024000,  # Default 1MB
        file_url=file_url,
        storage_key=storage_key,
        tags=data.tags or [],
        category=data.category,
        uploaded_by=data.uploaded_by,
        uploaded_by_name=uploaded_by_name,
        uploaded_at=datetime.now(),
        view_count=0,
        last_accessed_at=None,
    )
    
    _mock_knowledge_docs[str(doc_id)] = doc
    logger.info(f"[Mock] Uploaded knowledge document: {doc.title}")

    # Persist metadata to DB
    doc_persisted = False
    try:
        db.execute(
            text(
                """
                INSERT INTO knowledge_document (
                    id, title, description, source, category, tags,
                    file_type, file_size, storage_key, file_url,
                    org_id, project_id, meeting_id, visibility,
                    created_at, updated_at
                )
                VALUES (
                    :id, :title, :description, :source, :category, :tags,
                    :file_type, :file_size, :storage_key, :file_url,
                    :org_id, :project_id, :meeting_id, :visibility,
                    now(), now()
                )
                ON CONFLICT (id) DO NOTHING
                """
            ),
            {
                "id": str(doc_id),
                "title": doc.title,
                "description": doc.description,
                "source": doc.source,
                "category": doc.category,
                "tags": doc.tags,
                "file_type": doc.file_type,
                "file_size": doc.file_size,
                "storage_key": doc.storage_key,
                "file_url": doc.file_url,
                "org_id": None,
                "project_id": str(data.project_id) if data.project_id else None,
                "meeting_id": str(data.meeting_id) if data.meeting_id else None,
                "visibility": None,
            },
        )
        db.commit()
        doc_persisted = True
    except Exception as exc:
        logger.error("Failed to persist knowledge_document to DB: %s", exc)

    # Auto-embed and store chunks if HF is available
    try:
        if not doc_persisted:
            raise RuntimeError("knowledge_document row not persisted; skip embedding to avoid FK errors")
        text_content = ""
        # If original file content is available, try decode
        try:
            if file and "content" in locals() and content:
                # Extract by type
                if file_ext == "pdf":
                    text_content = _sanitize_text(_extract_pdf_text(content))
                elif file_ext in ("doc", "docx"):
                    text_content = _sanitize_text(_extract_docx_text(content))
                elif file_ext in ("xls", "xlsx"):
                    text_content = _sanitize_text(_extract_excel_text(content))
                else:
                    text_content = _sanitize_text(content.decode("utf-8", errors="ignore"))
        except Exception:
            text_content = ""

        if not text_content:
            # Fallback: use description/title
            text_parts = [doc.title]
            if doc.description:
                text_parts.append(doc.description)
            text_content = _sanitize_text("\n".join(text_parts))

        chunks = _chunk_text(text_content) if text_content else []
        if chunks and is_jina_available():
            chunks = [_sanitize_text(c) for c in chunks if _sanitize_text(c)]
            # skip if still empty
            if not chunks:
                raise ValueError("No valid text chunks to embed (maybe binary/pdf without extraction)")
            embed_inputs = [_normalize_for_embedding(c) for c in chunks]
            embeddings = embed_texts(embed_inputs)
            for idx, (chunk, emb) in enumerate(zip(chunks, embeddings)):
                emb_literal = "[" + ",".join(str(x) for x in emb) + "]"
                db.execute(
                    text(
                        """
                        INSERT INTO knowledge_chunk (
                            id, document_id, chunk_index, content, embedding, scope_meeting, scope_project, created_at
                        )
                        VALUES (
                            :id, :document_id, :chunk_index, :content, CAST(:embedding AS vector),
                            :scope_meeting, :scope_project, now()
                        )
                        """
                    ),
                    {
                        "id": str(uuid4()),
                        "document_id": str(doc_id),
                        "chunk_index": idx,
                        "content": chunk,
                        "embedding": emb_literal,
                        "scope_meeting": str(data.meeting_id) if getattr(data, "meeting_id", None) else None,
                        "scope_project": str(data.project_id) if getattr(data, "project_id", None) else None,
                    },
                )
            db.commit()
    except Exception as exc:
        logger.error("Auto-embed failed: %s", exc, exc_info=True)
    
    return KnowledgeDocumentUploadResponse(
        id=doc_id,
        title=doc.title,
        file_url=file_url,
        message="Tài liệu đã được tải lên thành công",
    )


async def ingest_document(db: Session, document_id: UUID) -> Optional[dict]:
    """
    Ingest a document into vector store.
    Current implementation is a stub: combines metadata text and upserts to PgVectorClient.
    """
    doc = _mock_knowledge_docs.get(str(document_id))
    if not doc:
        return None

    parts = [doc.title]
    if doc.description:
        parts.append(doc.description)
    if doc.tags:
        parts.append(" ".join(doc.tags))
    content = "\n".join(parts)

    settings = get_settings()
    client = PgVectorClient(connection=settings.database_url)
    client.upsert([content])

    logger.info("Ingested document %s into vector store (stub)", document_id)
    return {"status": "embedded", "document_id": document_id}


async def update_document(
    db: Session,
    document_id: UUID,
    data: KnowledgeDocumentUpdate,
) -> Optional[KnowledgeDocument]:
    """Update a document (DB first, fallback mock)"""
    update_data = data.model_dump(exclude_unset=True)
    # Try DB
    try:
        params = {"id": str(document_id)}
        fields = []
        for field in ["title", "description", "category", "tags", "meeting_id", "project_id"]:
            if field in update_data and update_data[field] is not None:
                params[field] = update_data[field]
                fields.append(f"{field} = :{field}")
        if fields:
            set_clause = ", ".join(fields + ["updated_at = now()"])
            row = db.execute(
                text(
                    f"""
                    UPDATE knowledge_document
                    SET {set_clause}
                    WHERE id = :id
                    RETURNING id, title, description, source, category, tags,
                              file_type, file_size, storage_key, file_url,
                              created_at, updated_at, NULL::text AS document_type
                    """
                ),
                params,
            ).mappings().first()
            if row:
                db.commit()
                return _with_presigned_url(_row_to_doc(row))
    except Exception as exc:
        logger.warning("DB update_document failed, fallback to mock: %s", exc)

    # Fallback mock
    doc = _mock_knowledge_docs.get(str(document_id))
    if not doc:
        return None
    
    doc_dict = doc.model_dump()
    doc_dict.update(update_data)
    
    updated_doc = KnowledgeDocument(**doc_dict)
    _mock_knowledge_docs[str(document_id)] = updated_doc
    return updated_doc


async def delete_document(db: Session, document_id: UUID) -> bool:
    """Delete a document: DB metadata, chunks, and stored file."""
    deleted = False
    # Fetch storage key and file_url for cleanup
    storage_key = None
    file_url = None
    try:
        row = db.execute(
            text("SELECT storage_key, file_url FROM knowledge_document WHERE id = :id"),
            {"id": str(document_id)},
        ).mappings().first()
        if row:
            storage_key = row.get("storage_key")
            file_url = row.get("file_url")
    except Exception as exc:
        logger.warning("Failed to fetch document before delete: %s", exc)

    try:
        # Delete chunks
        db.execute(
            text("DELETE FROM knowledge_chunk WHERE document_id = :id"),
            {"id": str(document_id)},
        )
        # Delete document
        result = db.execute(
            text("DELETE FROM knowledge_document WHERE id = :id"),
            {"id": str(document_id)},
        )
        db.commit()
        deleted = result.rowcount > 0
    except Exception as exc:
        logger.error("Failed to delete document %s: %s", document_id, exc, exc_info=True)
        db.rollback()
        deleted = False

    # Remove from mock cache
    if str(document_id) in _mock_knowledge_docs:
        del _mock_knowledge_docs[str(document_id)]

    # Delete file from storage if any
    try:
        if storage_key and is_storage_configured():
            delete_object(storage_key)
        # Delete local file if stored locally (/files/ or uploaded_files)
        if file_url and file_url.startswith("/files/"):
            local_path = Path(__file__).parent.parent / file_url.lstrip("/")
            if local_path.exists():
                local_path.unlink()
    except Exception as exc:
        logger.warning("Cleanup storage for %s failed: %s", document_id, exc)

    return deleted


def _build_vector_filters(request: KnowledgeSearchRequest):
    filters = ["1=1"]
    params = {
        "chunk_limit": max(request.limit * 4, 20),
    }
    if request.source:
        filters.append("kd.source = :source")
        params["source"] = request.source
    if request.category:
        filters.append("kd.category = :category")
        params["category"] = request.category
    # Tags filter kept simple: intersects
    if request.tags:
        filters.append("kd.tags && :tags")
        params["tags"] = request.tags
    if getattr(request, "meeting_id", None):
        filters.append("COALESCE(kd.meeting_id, kc.scope_meeting) = :meeting_id")
        params["meeting_id"] = str(request.meeting_id)
    if getattr(request, "project_id", None):
        filters.append("COALESCE(kd.project_id, kc.scope_project) = :project_id")
        params["project_id"] = str(request.project_id)
    return " AND ".join(filters), params


async def _vector_search_documents(
    db: Session,
    request: KnowledgeSearchRequest,
) -> Optional[KnowledgeSearchResponse]:
    """Search using pgvector; returns None on failure."""
    if not is_jina_available():
        return None
    try:
        query_vec = embed_texts([_normalize_for_embedding(request.query)])[0]
        vec_literal = _format_vector(query_vec)

        where_clause, params = _build_vector_filters(request)
        params.update(
            {
                "query_vec": vec_literal,
                "offset": request.offset,
            }
        )

        rows = db.execute(
            text(
                f"""
                SELECT
                    kd.id,
                    kd.title,
                    kd.description,
                    kd.source,
                    kd.category,
                    kd.tags,
                    kd.file_type,
                    kd.file_size,
                    kd.storage_key,
                    kd.file_url,
                    kd.created_at,
                    kd.updated_at,
                    NULL::text AS document_type,
                    (kc.embedding <=> CAST(:query_vec AS vector)) AS distance
                FROM knowledge_chunk kc
                JOIN knowledge_document kd ON kc.document_id = kd.id
                WHERE {where_clause}
                ORDER BY distance ASC
                LIMIT :chunk_limit
                """
            ),
            params,
        ).mappings().all()

        if not rows:
            return None

        # Deduplicate by document, keep best distance
        doc_best = {}
        for r in rows:
            doc_id = r["id"]
            dist = r["distance"]
            if doc_id not in doc_best or dist < doc_best[doc_id]["distance"]:
                doc_best[doc_id] = {"row": r, "distance": dist}

        ordered = sorted(doc_best.values(), key=lambda x: x["distance"])
        ordered = ordered[request.offset : request.offset + request.limit]

        docs = [_with_presigned_url(_row_to_doc(item["row"])) for item in ordered]
        total = len(doc_best)
        return KnowledgeSearchResponse(documents=docs, total=total, query=request.query)
    except Exception as exc:
        logger.error("Vector search failed, fallback to text: %s", exc, exc_info=True)
        return None


async def search_documents(
    db: Session,
    request: KnowledgeSearchRequest,
) -> KnowledgeSearchResponse:
    """Vector search first; fallback to simple text search."""
    # Vector search using Jina embedding + pgvector
    vector_result = await _vector_search_documents(db, request)
    if vector_result:
        return vector_result

    # Fallback to text search
    try:
        conditions = ["1=1"]
        params = {
            "like_q": f"%{request.query}%",
            "offset": request.offset,
            "limit": request.limit,
        }
        conditions.append("(title ILIKE :like_q OR description ILIKE :like_q OR :like_q = ANY(tags))")
        if request.source:
            conditions.append("source = :source")
            params["source"] = request.source
        if request.category:
            conditions.append("category = :category")
            params["category"] = request.category
        if request.meeting_id:
            conditions.append("meeting_id = :meeting_id")
            params["meeting_id"] = str(request.meeting_id)
        if request.project_id:
            conditions.append("project_id = :project_id")
            params["project_id"] = str(request.project_id)

        where_clause = " AND ".join(conditions)
        rows = db.execute(
            text(
                f"""
                SELECT id, title, description, source, category, tags,
                       file_type, file_size, storage_key, file_url,
                       created_at, updated_at,
                       NULL::text AS document_type
                FROM knowledge_document
                WHERE {where_clause}
                ORDER BY created_at DESC NULLS LAST
                LIMIT :limit OFFSET :offset
                """
            ),
            params,
        ).mappings().all()

        total = db.execute(
            text(f"SELECT COUNT(*) FROM knowledge_document WHERE {where_clause}"),
            params,
        ).scalar_one()

        docs = [_with_presigned_url(_row_to_doc(r)) for r in rows]
        return KnowledgeSearchResponse(documents=docs, total=total, query=request.query)
    except Exception as exc:
        logger.warning("Search documents failed, returning empty: %s", exc)
        return KnowledgeSearchResponse(documents=[], total=0, query=request.query)


async def query_knowledge_ai(
    db: Session,
    request: KnowledgeQueryRequest,
) -> KnowledgeQueryResponse:
    """RAG query using pgvector + Groq"""
    # Smalltalk/noise handling
    if _is_smalltalk_or_noise(request.query):
        answer = "Xin chào! Bạn muốn hỏi gì về tài liệu/policy? Hãy mô tả rõ hơn nhé."
        return KnowledgeQueryResponse(
            answer=answer,
            relevant_documents=[],
            confidence=0.5,
            citations=[],
        )

    top_k_chunks = max(request.limit * 3, 12)
    chunks = []
    relevant_docs: List[KnowledgeDocument] = []
    citations: List[str] = []
    best_score = None

    if is_jina_available():
        try:
            query_vec = embed_texts([_normalize_for_embedding(request.query)])[0]
            vec_literal = _format_vector(query_vec)

            where_clause, params = _build_vector_filters(
                KnowledgeSearchRequest(
                    query=request.query,
                    limit=top_k_chunks,
                    offset=0,
                    source=None,
                    category=None,
                    tags=None,
                    meeting_id=request.meeting_id,
                    project_id=request.project_id,
                )
            )
            params.update({"query_vec": vec_literal, "top_k": top_k_chunks})

            rows = db.execute(
                text(
                    f"""
                    SELECT
                        kd.id,
                        kd.title,
                        kd.description,
                        kd.source,
                        kd.category,
                        kd.tags,
                        kd.file_type,
                    kd.file_size,
                    kd.storage_key,
                    kd.file_url,
                    kd.created_at,
                    kd.updated_at,
                    NULL::text AS document_type,
                    kc.content,
                    kc.chunk_index,
                    (kc.embedding <=> CAST(:query_vec AS vector)) AS distance
                FROM knowledge_chunk kc
                JOIN knowledge_document kd ON kc.document_id = kd.id
                    WHERE {where_clause}
                    ORDER BY distance ASC
                    LIMIT :top_k
                    """
                ),
                params,
            ).mappings().all()

            # Dedup docs and collect top chunks
            doc_best = {}
            for r in rows:
                doc_id = r["id"]
                dist = r["distance"]
                if doc_id not in doc_best or dist < doc_best[doc_id]["distance"]:
                    doc_best[doc_id] = {"row": r, "distance": dist}
                chunks.append(
                    {
                        "doc_id": doc_id,
                        "title": r["title"],
                        "distance": dist,
                        "text": _sanitize_text(r["content"])[:800],
                    }
                )

            ordered_docs = sorted(doc_best.values(), key=lambda x: x["distance"])
            relevant_docs = [_with_presigned_url(_row_to_doc(item["row"])) for item in ordered_docs[: request.limit]]
            citations = [d.title for d in relevant_docs]
            best_score = ordered_docs[0]["distance"] if ordered_docs else None
        except Exception as exc:
            logger.error("RAG query vector path failed: %s", exc, exc_info=True)

    # Build context
    context_parts = []
    for ch in chunks[: top_k_chunks]:
        context_parts.append(f"[{ch['title']} | score={ch['distance']:.3f}] {ch['text']}")
    context = "\n".join(context_parts) if context_parts else "Không có ngữ cảnh liên quan."

    # If no context at all, avoid repeating 'Không đủ dữ liệu', give gentle ask for clarification
    if not context_parts:
        # Try giving a gentle, generic answer using Groq (no RAG context)
        if is_gemini_available():
            try:
                chat = GeminiChat(
                    system_prompt=(
                        "Bạn là trợ lý thân thiện. Trả lời ngắn gọn, tiếng Việt. "
                        "Không bịa quá đà; nếu không chắc, hãy nói rõ."
                    )
                )
                prompt = f"""Câu hỏi: {request.query}
Hiện chưa tìm thấy tài liệu phù hợp trong hệ thống.
Hãy:
- Nói rõ chưa có tài liệu khớp và đề nghị người dùng mô tả thêm.
- Sau đó đưa ra gợi ý chung (mang tính kiến thức nền, có thể không chính xác tuyệt đối)."""
                answer = await chat.chat(prompt)
                return KnowledgeQueryResponse(
                    answer=answer,
                    relevant_documents=[],
                    confidence=0.35,
                    citations=[],
                )
            except Exception as exc:
                logger.error("Groq generic fallback failed: %s", exc)
        return KnowledgeQueryResponse(
            answer="Mình chưa thấy tài liệu liên quan. Bạn mô tả rõ hơn chủ đề/tên tài liệu nhé? Nếu cần, mình có thể gợi ý chung.",
            relevant_documents=[],
            confidence=0.3,
            citations=[],
        )

    # Call Groq LLM
    if is_gemini_available():
        try:
            chat = GeminiChat(
                system_prompt=(
                    "Bạn là trợ lý RAG. Trả lời ngắn gọn bằng tiếng Việt. "
                    "Chỉ dùng thông tin trong Context. Nếu thiếu thông tin, nói rõ."
                )
            )
            prompt = f"""Câu hỏi: {request.query}

Context (top chunks):
{context}

Yêu cầu:
- Trả lời ngắn gọn, không markdown.
- Nếu dùng thông tin, nêu rõ tên tài liệu trong ngoặc [].
- Nếu không đủ thông tin, trả lời rằng không đủ dữ liệu."""
            answer = await chat.chat(prompt)
            confidence = 0.90 if relevant_docs else 0.60
            if best_score is not None:
                confidence = max(0.5, min(0.98, 1 - float(best_score)))
            return KnowledgeQueryResponse(
                answer=answer,
                relevant_documents=relevant_docs,
                confidence=confidence,
                citations=citations,
            )
        except Exception as exc:
            logger.error("Groq query failed: %s", exc)

    # Fallback response
    if relevant_docs:
        answer = f"Tìm thấy {len(relevant_docs)} tài liệu liên quan: {', '.join([d.title for d in relevant_docs[:3]])}"
    else:
        answer = "Không tìm thấy thông tin liên quan trong knowledge base."

    return KnowledgeQueryResponse(
        answer=answer,
        relevant_documents=relevant_docs,
        confidence=0.70,
        citations=citations,
    )
